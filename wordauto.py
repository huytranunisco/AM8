from docx import Document
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import Select
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver import ActionChains
from pyautogui import screenshot
from PIL import Image
import os
import time
from report_automate.BNPauto import facilityMatcher

def billPeriodMenu():
    menu = 'Billing Period for Handling:\n1. Bimonthly\n2. Weekly\n3. Daily\n'
    period = input(menu)
    while period != '1' and period != '2' and period != '3':
        period = input('Choose one of the options (1-3).')
    
    options = ['bimonthly','weekly', 'daily']

    period = options[int(period) - 1]

    return period

def replaceString(paragraph, oldText, newText):
    inline = paragraph.runs
    for i in range(len(inline)):
        if oldText in inline[i].text:
            text = inline[i].text.replace(oldText, newText)
            inline[i].text = text
        if " (1-15, 16-EOM)" in inline[i].text and billingPeriod != 'bimonthly':
            text = inline[i].text.replace(' (1-15, 16-EOM)', '')
            inline[i].text = text
    
def imageCapture(accName, accFacility, directory):
    accName = accName[:5]
    facility = facilityMatcher(accFacility)
    
    #Getting BNP
    driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()))
    driver.get("http://bnp.unisco.com/")
    action = ActionChains(driver)
    assert len(driver.window_handles) == 1

    driver.maximize_window()

    #Logging into BNP
    interactor = driver.find_element(By.ID,"inputUserName")
    interactor.send_keys('wiserpa')
    interactor = driver.find_element(By.ID,"inputPassword")
    interactor.send_keys('#rpa#1234')
    interactor = driver.find_element(By.XPATH,"/html/body/div/footer/div/button")
    interactor.click()

    #Clicking Sales Module from Module Dropdown Menu
    select = WebDriverWait(driver, 10).until(EC.element_to_be_clickable((By.ID, 'TS_span_menu')))
    select.click()
    select = WebDriverWait(driver, 5).until(EC.presence_of_element_located((By.XPATH, '//*[@id="headmenu_mn_active"]/div/ul/li[1]')))
    select.click()

    #Clicking Invoice Management from Invoice Dropdown Menu
    select = WebDriverWait(driver, 5).until(EC.presence_of_element_located((By.XPATH, '//*[@id="headmenu"]/li[3]/span')))
    select.click()
    select = WebDriverWait(driver, 5).until(EC.presence_of_element_located((By.XPATH, '/html/body/div/div/header/div[1]/ul/li[3]/div/ul/li[1]/a')))
    select.click()

    #Inputting information on the Invoice Management Page and Searching

    #Bill To
    interactor = driver.find_element(By.XPATH, '//*[@id="sitecontent"]/div[2]/div[1]/div[6]/span/span/input')
    interactor.send_keys(accName)
    time.sleep(0.5)
    interactor.send_keys(Keys.ENTER)

    #Facility
    select = Select(driver.find_element(By.ID, 'ddlFacility'))
    select.select_by_visible_text(facility)

    #Invoice Status
    interactor = driver.find_element(By.XPATH, '//*[@id="sitecontent"]/div[2]/div[2]/div[5]/div/div/input')
    interactor.send_keys('Check')
    time.sleep(0.5)
    interactor.send_keys(Keys.ENTER)

    #Period Start
    interactor = driver.find_element(By.ID, 'inputPeriodStart')
    interactor.send_keys('08/01/22')

    #Period End
    interactor = driver.find_element(By.ID, 'inputPeriodEnd')
    interactor.send_keys('08/15/22')

    #Category
    interactor = driver.find_element(By.XPATH, '//*[@id="sitecontent"]/div[2]/div[2]/div[6]/div/input')
    interactor.send_keys('Handling')

    #Clicking Search
    interactor = driver.find_element(By.XPATH, '/html/body/div[1]/div/div/div[2]/div[3]/div[11]/button')
    interactor.click()
    interactor = driver.find_element(By.XPATH, '/html/body/div[1]/div/div/div[2]/div[3]/div[11]/button')
    interactor.click()
    time.sleep(2)

    imagescreen = screenshot()
    imagescreen.save(directory + '\\image1.png')
    print('First image saved!')

    #Checking first invoice
    interactor = driver.find_element(By.XPATH, '/html/body/div[1]/div/div/div[3]/div/div/div[3]/table/tbody/tr/td[2]/div[1]/a[1]')
    action.move_to_element(interactor).perform()
    interactor.click()

    time.sleep(4)
    imagescreen = screenshot()
    imagescreen.save(directory + '\\image2.png')
    print('Second image saved!')

    driver.switch_to.window(driver.window_handles[0])
    driver.get('https://wise.logisticsteam.com/v2/#/login')

    driver.maximize_window()

    interactor = WebDriverWait(driver, 15).until(EC.presence_of_element_located((By.XPATH, '/html/body/div1/div/div[2]/div/div/div/form/input[1]')))
    interactor.send_keys('marionz')
    interactor = driver.find_element(By.NAME,"password")
    interactor.send_keys('qwer1234')
    interactor = WebDriverWait(driver, 5).until(EC.presence_of_element_located((By.XPATH, '//*[@id="loginBtn"]/button')))
    interactor.click()

    interactor = WebDriverWait(driver, 5).until(EC.presence_of_element_located((By.XPATH, '/html/body/div1/header/div[1]/div[5]/ul/li[1]/ul')))
    items = interactor.find_elements(By.TAG_NAME,'li')
    for index, company in enumerate(items):
        try:
            company = company.find_element(By.TAG_NAME,'a')
            item = company.find_elements(By.TAG_NAME,'div')[1]
            itemtext = item.get_attribute('textContent')
            if facility in itemtext:
                company.click()
        except:
            continue

    interactor = driver.find_element(By.XPATH, '/html/body/div1/header/div[1]/div[5]/ul/li[1]/span')
    interactor.click()
    
    driver.switch_to.window(driver.window_handles[0])
    time.sleep(3)

    imagescreen = screenshot()
    imagescreen.save(directory + '\\image3.png')
    print('Third image saved!')

def imageCrop(path):
    image_one = Image.open(path + '\\image1.png')
    image_one = image_one.crop((10, 410, 1370, 550))
    image_one.save(path + '\\image1.png')
    
    image_two = Image.open(path + '\\image2.png')
    image_two = image_two.crop((15, 175, 1760, 700))
    image_two.save(path + '\\image2.png')

    image_three = Image.open(path + '\\image3.png')
    image_three = image_three.crop((1310, 120, 1900, 615))
    image_three.save(path + '\\image3.png')

accName = input('Enter Name of Customer (No commas): ')
accFacility = input('Enter Facility of Customer: ')
billerName = input("Enter Biller's Name for Customer: ")
userPath ='kenguyen'
billingPeriod = billPeriodMenu()

directory = f'{accName}-{accFacility}'
parentDir = "C:\\Users\\" + userPath + "\\Desktop\\"
path = os.path.join(parentDir, directory)
try:
    os.mkdir(path)
    print(f'Directory "{directory}" created')
except (FileExistsError):
    print('Folder already exists. Moving on...')

imageCapture(accName, accFacility, path)
imageCrop(path)

tempFac = accFacility.replace(" ", "")
tempAcc = accName.replace(" ", "")

oldText = {"One" : f'{accName} ({accFacility}) SOP', "Biller Name" : billerName, "billingPeriodH" : billingPeriod, "Two" : f'{accName} ({accFacility})',
           "Three" : f'{tempAcc}-{tempFac}-{billingPeriod.capitalize()}', "accName" : accName, "billingfac" : accFacility, "billingperiodh" : billingPeriod}
oldTextList = oldText.keys()

document = Document("C:\\Users\\kenguyen\\Documents\\SOPS\\SOP Template.docx")

for index, paragraph in enumerate(document.paragraphs):
    if index == 39:
        paragraph.add_run().add_picture(path + '\\image1.png')
    elif index == 41:
        paragraph.add_run().add_picture(path + '\\image2.png')
    elif index == 61:
        paragraph.add_run().add_picture(path + '\\image3.png')
    for text in oldTextList:
        if text in paragraph.text:
            replaceString(paragraph, text, oldText[text])

textChange = f'{tempAcc}-{tempFac}-{billingPeriod.capitalize()}-SOP.docx'
document.save(path + '\\' + textChange)
print('Saved new document!')