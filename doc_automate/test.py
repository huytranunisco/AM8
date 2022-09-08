from argparse import Action
from msilib.schema import File
from docx import Document
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import Select
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver import ActionChains
from pyautogui import screenshot
from PIL import Image
import os
import time
from BNPauto import facilityMatcher

def billPeriodMenu():
    menu = 'Billing Period for Handling:\n1. Bimonthly\n2. Weekly\n3. Daily\n'
    period = input(menu)
    while period != '1' and period != '2' and period != '3':
        period = input('Choose one of the options (1-3).')
    
    options = ['bimonthly','weekly', 'daily']

    period = options[int(period) - 1]

    return period

def replaceString(paragraph, oldText, newText):
    inline = paragraph.runs
    for i in range(len(inline)):
        if oldText in inline[i].text:
            text = inline[i].text.replace(oldText, newText)
            inline[i].text = text
            break

def imageCapture(accName, accFacility, directory):
    accName = accName[:5]
    facility = facilityMatcher(accFacility)

    #Getting BNP
    driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()))
    driver.get("http://bnp.unisco.com/")
    action = ActionChains(driver)
    assert len(driver.window_handles) == 1

    driver.maximize_window()

    #Logging into BNP
    interactor = driver.find_element(By.ID,"inputUserName")
    interactor.send_keys('wiserpa')
    interactor = driver.find_element(By.ID,"inputPassword")
    interactor.send_keys('#rpa#1234')
    interactor = driver.find_element(By.XPATH,"/html/body/div/footer/div/button")
    interactor.click()

    #Clicking Sales Module from Module Dropdown Menu
    select = WebDriverWait(driver, 5).until(EC.element_to_be_clickable((By.ID, 'TS_span_menu')))
    select.click()
    select = WebDriverWait(driver, 5).until(EC.presence_of_element_located((By.XPATH, '//*[@id="headmenu_mn_active"]/div/ul/li[1]')))
    select.click()

    #Clicking Invoice Management from Invoice Dropdown Menu
    select = WebDriverWait(driver, 5).until(EC.presence_of_element_located((By.XPATH, '//*[@id="headmenu"]/li[3]/span')))
    select.click()
    select = WebDriverWait(driver, 5).until(EC.presence_of_element_located((By.XPATH, '/html/body/div/div/header/div[1]/ul/li[3]/div/ul/li[1]/a')))
    select.click()

    #Inputting information on the Invoice Management Page and Searching

    #Bill To
    interactor = driver.find_element(By.XPATH, '//*[@id="sitecontent"]/div[2]/div[1]/div[6]/span/span/input')
    interactor.send_keys(accName)
    time.sleep(0.5)
    interactor.send_keys(Keys.ENTER)

    #Facility
    select = Select(driver.find_element(By.ID, 'ddlFacility'))
    select.select_by_visible_text(facility)

    #Invoice Status
    interactor = driver.find_element(By.XPATH, '//*[@id="sitecontent"]/div[2]/div[2]/div[5]/div/div/input')
    interactor.send_keys('Check')
    time.sleep(0.5)
    interactor.send_keys(Keys.ENTER)

    #Period Start
    interactor = driver.find_element(By.ID, 'inputPeriodStart')
    interactor.send_keys('08/01/22')

    #Period End
    interactor = driver.find_element(By.ID, 'inputPeriodEnd')
    interactor.send_keys('08/15/22')

    #Category
    interactor = driver.find_element(By.XPATH, '//*[@id="sitecontent"]/div[2]/div[2]/div[6]/div/input')
    interactor.send_keys('Handling')

    #Clicking Search
    interactor = driver.find_element(By.XPATH, '/html/body/div[1]/div/div/div[2]/div[3]/div[11]/button')
    interactor.click()
    interactor = driver.find_element(By.XPATH, '/html/body/div[1]/div/div/div[2]/div[3]/div[11]/button')
    interactor.click()
    time.sleep(2)

    imagescreen = screenshot()
    imagescreen.save(directory + '\\image1.png')
    print('First image saved!')

    #Checking first invoice
    interactor = driver.find_element(By.XPATH, '/html/body/div[1]/div/div/div[3]/div/div/div[3]/table/tbody/tr/td[2]/div[1]/a[1]')
    action.move_to_element(interactor).perform()
    interactor.click()

    imagescreen = screenshot()
    imagescreen.save(directory + '\\image2.png')
    print('Second image saved!')




accName = 'E&S International Enterprises Inc'
accFacility = 'Valley View'
billerName = 'Luke Darang'
userPath = input('Username for PC: ')
billingPeriod = billPeriodMenu()

directory = f'{accName}-{accFacility}'
parentDir = "C:\\Users\\" + userPath + "\\Desktop\\"
path = os.path.join(parentDir, directory)
try:
    os.mkdir(path)
    print(f'Directory "{directory}" created')
except (FileExistsError):
    print('File already exists. Moving on...')

imageCapture(accName, accFacility, path)

tempFac = accFacility.replace(" ", "")
tempAcc = accName.replace(" ", "")

oldText = {"One" : f'{accName} ({accFacility}) SOP', "Biller Name" : billerName, "billingPeriodH" : billingPeriod, "Two" : f'{accName} ({accFacility})',
           "Three" : f'{tempAcc}-{tempFac}-{billingPeriod.capitalize()}', "accName" : accName, "Fac" : accFacility, "billingperiodh" : billingPeriod}
oldTextList = oldText.keys()

document = Document("C:\\Users\\kenguyen\\Documents\\SOPS\\SOP Template.docx")

if billingPeriod != 'bimonthly':
    document.paragraphs[3].text = document.paragraphs[3].text.replace(" (1-15, 16-EOM)", "")
    document.paragraphs[4].text = document.paragraphs[4].text.replace(" (1-15, 16-EOM)", "")

for paragraph in document.paragraphs:
    for text in oldTextList:
        if text in paragraph.text:
            replaceString(paragraph, text, oldText[text])

textChange = f'{tempAcc}-{tempFac}-{billingPeriod.capitalize()}-SOP.docx'
document.save(textChange)
